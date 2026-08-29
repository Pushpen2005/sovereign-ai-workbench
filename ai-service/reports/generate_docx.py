#!/usr/bin/env python3
"""
PR #16 — Approval Note DOCX Generator

Reads sanitized JSON payload from stdin and generates a formatted,
executive-ready DOCX document using python-docx.
"""

import sys
import json
import os
import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def sanitize_value(val):
    if val is None:
        return "N/A"
    s = str(val).strip()
    if s.lower() in ("none", "null", "undefined", ""):
        return "N/A"
    return s

def set_cell_background(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_with_spacing(doc, text, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(26, 54, 93)  # Navy
    elif level == 2:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(45, 55, 72)
    return p

def create_document(data, output_path):
    doc = Document()

    # Configure Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Header
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = "SovereignAI — Industrial Inspection & Approval Note"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if header_p.runs:
            header_p.runs[0].font.name = "Arial"
            header_p.runs[0].font.size = Pt(8.5)
            header_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = "CONFIDENTIAL — For Authorized Operational Use Only"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_p.runs:
            footer_p.runs[0].font.name = "Arial"
            footer_p.runs[0].font.size = Pt(8.5)
            footer_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # ─── Title ───────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(18)
    title_run = title_p.add_run("APPROVAL NOTE")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 54, 93)

    # Subtitle / Document Type
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(16)
    sub_run = sub_p.add_run("Automated Operational Assessment & Executive Review")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)

    # ─── Section 1: Subject ──────────────────────────────────────────────────
    add_heading_with_spacing(doc, "1. Subject", level=1, space_before=10)
    subject_text = data.get("subject") or "Inspection Report Analysis and Approval Recommendation"
    p_subj = doc.add_paragraph()
    p_subj.paragraph_format.space_after = Pt(8)
    r_subj = p_subj.add_run(sanitize_value(subject_text))
    r_subj.font.name = "Arial"
    r_subj.font.size = Pt(10.5)

    # ─── Section 2: Background ───────────────────────────────────────────────
    add_heading_with_spacing(doc, "2. Background", level=1)
    background_text = data.get("background")
    if not background_text or sanitize_value(background_text) == "N/A":
        # Safe default summary derived strictly from provided findings
        findings = data.get("findings") or []
        eq_names = [f.get("equipment") for f in findings if f.get("equipment")]
        eq_summary = f"on {', '.join(set(eq_names))}" if eq_names else ""
        background_text = (
            f"An industrial inspection identified {len(findings)} finding(s) {eq_summary}. "
            "Operational parameters were recorded during standard inspection routines and evaluated "
            "against documented standard operating procedures (SOPs) to produce this assessment."
        ).strip()

    p_bg = doc.add_paragraph()
    p_bg.paragraph_format.space_after = Pt(8)
    r_bg = p_bg.add_run(sanitize_value(background_text))
    r_bg.font.name = "Arial"
    r_bg.font.size = Pt(10)

    # ─── Section 3: Inspection Findings ──────────────────────────────────────
    add_heading_with_spacing(doc, "3. Inspection Findings", level=1)
    findings = data.get("findings") or []

    if not findings:
        p_nf = doc.add_paragraph()
        p_nf.paragraph_format.space_after = Pt(8)
        r_nf = p_nf.add_run("No inspection findings recorded.")
        r_nf.font.name = "Arial"
        r_nf.font.italic = True
    else:
        for idx, finding in enumerate(findings, start=1):
            if len(findings) > 1:
                add_heading_with_spacing(doc, f"Finding {idx}", level=2, space_before=6, space_after=4)

            table = doc.add_table(rows=6, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            col_widths = [Inches(1.8), Inches(4.7)]
            fields = [
                ("Finding", sanitize_value(finding.get("finding"))),
                ("Equipment", sanitize_value(finding.get("equipment"))),
                ("Observed Value", sanitize_value(finding.get("observedValue"))),
                ("Limit", sanitize_value(finding.get("limit"))),
                ("Severity", sanitize_value(finding.get("severity"))),
                ("Evidence", sanitize_value(finding.get("evidence"))),
            ]

            for r_idx, (label, val) in enumerate(fields):
                row = table.rows[r_idx]

                # Label cell
                c0 = row.cells[0]
                c0.width = col_widths[0]
                set_cell_background(c0, "EDF2F7")  # light gray
                set_cell_margins(c0, top=70, bottom=70, left=100, right=100)
                p0 = c0.paragraphs[0]
                p0.paragraph_format.space_before = Pt(2)
                p0.paragraph_format.space_after = Pt(2)
                r0 = p0.add_run(label)
                r0.font.name = "Arial"
                r0.font.bold = True
                r0.font.size = Pt(9.5)
                r0.font.color.rgb = RGBColor(45, 55, 72)

                # Value cell
                c1 = row.cells[1]
                c1.width = col_widths[1]
                set_cell_background(c1, "FFFFFF")
                set_cell_margins(c1, top=70, bottom=70, left=100, right=100)
                p1 = c1.paragraphs[0]
                p1.paragraph_format.space_before = Pt(2)
                p1.paragraph_format.space_after = Pt(2)
                r1 = p1.add_run(val)
                r1.font.name = "Arial"
                r1.font.size = Pt(9.5)

            # Space after table
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)

    # ─── Section 4: Technical Analysis ───────────────────────────────────────
    add_heading_with_spacing(doc, "4. Technical Analysis", level=1)
    tech_analysis = data.get("technicalAnalysis")
    citations = data.get("citations") or []
    risk_assessment = data.get("riskAssessment") or {}

    if not tech_analysis or sanitize_value(tech_analysis) == "N/A":
        if citations:
            doc_names = list({c.get("filename") for c in citations if c.get("filename")})
            doc_str = ", ".join(doc_names) if doc_names else "authoritative SOPs"
            reason = risk_assessment.get("reason", "")
            tech_analysis = (
                f"Evaluation was performed by correlating observed parameters against standard operating procedures ({doc_str}). "
                + (f"Assessment analysis indicates: {reason}" if reason else "")
            ).strip()
        else:
            tech_analysis = "Technical analysis is unavailable based on the supplied SOP evidence."

    p_ta = doc.add_paragraph()
    p_ta.paragraph_format.space_after = Pt(8)
    r_ta = p_ta.add_run(sanitize_value(tech_analysis))
    r_ta.font.name = "Arial"
    r_ta.font.size = Pt(10)

    # ─── Section 5: Risk Assessment ──────────────────────────────────────────
    add_heading_with_spacing(doc, "5. Risk Assessment", level=1)
    risk_level = risk_assessment.get("level")
    risk_level_str = sanitize_value(risk_level).upper() if risk_level is not None else "Not Determined"
    risk_reason = sanitize_value(risk_assessment.get("reason"))

    risk_table = doc.add_table(rows=2, cols=2)
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    risk_table.autofit = False

    r_widths = [Inches(1.8), Inches(4.7)]
    for r in risk_table.rows:
        r.cells[0].width = r_widths[0]
        r.cells[1].width = r_widths[1]

    # Row 0: Risk Level
    c_rl_lbl = risk_table.rows[0].cells[0]
    set_cell_background(c_rl_lbl, "EDF2F7")
    set_cell_margins(c_rl_lbl, top=80, bottom=80, left=100, right=100)
    p_rl_lbl = c_rl_lbl.paragraphs[0]
    r_rl_lbl = p_rl_lbl.add_run("Risk Level")
    r_rl_lbl.font.name = "Arial"
    r_rl_lbl.font.bold = True
    r_rl_lbl.font.size = Pt(10)

    c_rl_val = risk_table.rows[0].cells[1]
    # Set background tint according to level
    if risk_level_str == "HIGH":
        set_cell_background(c_rl_val, "FED7D7")  # light red
    elif risk_level_str == "MEDIUM":
        set_cell_background(c_rl_val, "FEEBC8")  # light amber
    elif risk_level_str == "LOW":
        set_cell_background(c_rl_val, "C6F6D5")  # light green
    else:
        set_cell_background(c_rl_val, "F7FAFC")
    set_cell_margins(c_rl_val, top=80, bottom=80, left=100, right=100)

    p_rl_val = c_rl_val.paragraphs[0]
    r_rl_val = p_rl_val.add_run(risk_level_str)
    r_rl_val.font.name = "Arial"
    r_rl_val.font.bold = True
    r_rl_val.font.size = Pt(11)
    if risk_level_str == "HIGH":
        r_rl_val.font.color.rgb = RGBColor(197, 48, 48)
    elif risk_level_str == "MEDIUM":
        r_rl_val.font.color.rgb = RGBColor(183, 121, 31)
    elif risk_level_str == "LOW":
        r_rl_val.font.color.rgb = RGBColor(40, 116, 101)

    # Row 1: Reason
    c_rr_lbl = risk_table.rows[1].cells[0]
    set_cell_background(c_rr_lbl, "EDF2F7")
    set_cell_margins(c_rr_lbl, top=80, bottom=80, left=100, right=100)
    p_rr_lbl = c_rr_lbl.paragraphs[0]
    r_rr_lbl = p_rr_lbl.add_run("Reason")
    r_rr_lbl.font.name = "Arial"
    r_rr_lbl.font.bold = True
    r_rr_lbl.font.size = Pt(10)

    c_rr_val = risk_table.rows[1].cells[1]
    set_cell_background(c_rr_val, "FFFFFF")
    set_cell_margins(c_rr_val, top=80, bottom=80, left=100, right=100)
    p_rr_val = c_rr_val.paragraphs[0]
    r_rr_val = p_rr_val.add_run(risk_reason)
    r_rr_val.font.name = "Arial"
    r_rr_val.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ─── Section 6: Recommendation ───────────────────────────────────────────
    add_heading_with_spacing(doc, "6. Recommendation", level=1)
    recommendation_text = sanitize_value(data.get("recommendation"))

    rec_table = doc.add_table(rows=1, cols=1)
    rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rec_cell = rec_table.rows[0].cells[0]
    rec_cell.width = Inches(6.5)
    set_cell_background(rec_cell, "EBF8FF")  # subtle light blue
    set_cell_margins(rec_cell, top=100, bottom=100, left=120, right=120)

    p_rec = rec_cell.paragraphs[0]
    p_rec.paragraph_format.space_before = Pt(4)
    p_rec.paragraph_format.space_after = Pt(4)
    r_rec = p_rec.add_run(recommendation_text)
    r_rec.font.name = "Arial"
    r_rec.font.size = Pt(10)
    r_rec.font.bold = True
    r_rec.font.color.rgb = RGBColor(43, 108, 176)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ─── Section 7: References ───────────────────────────────────────────────
    add_heading_with_spacing(doc, "7. References", level=1)
    if not citations:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(8)
        r_ref = p_ref.add_run("No specific SOP references cited.")
        r_ref.font.name = "Arial"
        r_ref.font.italic = True
    else:
        for idx, citation in enumerate(citations, start=1):
            fn = sanitize_value(citation.get("filename"))
            pg = sanitize_value(citation.get("page"))
            ci = sanitize_value(citation.get("chunkIndex"))
            doc_id = citation.get("documentId")

            ref_parts = [f"Filename: {fn}", f"Page: {pg}"]
            if ci != "N/A":
                ref_parts.append(f"Chunk Index: {ci}")
            if doc_id and sanitize_value(doc_id) != "N/A":
                ref_parts.append(f"Document ID: {doc_id}")

            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.left_indent = Inches(0.25)
            p_ref.paragraph_format.space_after = Pt(4)
            r_num = p_ref.add_run(f"[{idx}] ")
            r_num.font.name = "Arial"
            r_num.font.bold = True
            r_num.font.size = Pt(9.5)
            r_text = p_ref.add_run(" | ".join(ref_parts))
            r_text.font.name = "Arial"
            r_text.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ─── Section 8: Approval ─────────────────────────────────────────────────
    add_heading_with_spacing(doc, "8. Approval", level=1)

    appr_table = doc.add_table(rows=4, cols=2)
    appr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    appr_table.autofit = False

    a_widths = [Inches(2.0), Inches(4.5)]
    appr_rows = [
        ("Prepared By:", "SovereignAI"),
        ("Reviewed By:", "____________________________________"),
        ("Approved By:", "____________________________________"),
        ("Date:", "____________________________________"),
    ]

    for idx, (label, val) in enumerate(appr_rows):
        row = appr_table.rows[idx]
        row.cells[0].width = a_widths[0]
        row.cells[1].width = a_widths[1]

        c0 = row.cells[0]
        set_cell_background(c0, "F7FAFC")
        set_cell_margins(c0, top=60, bottom=60, left=80, right=80)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        r0.font.name = "Arial"
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = RGBColor(74, 85, 104)

        c1 = row.cells[1]
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c1, top=60, bottom=60, left=80, right=80)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(val)
        r1.font.name = "Arial"
        r1.font.size = Pt(9.5)

    # Ensure target directory exists
    out_dir = os.path.dirname(os.path.abspath(output_path))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc.save(output_path)
    abs_path = os.path.abspath(output_path)
    return abs_path

def main():
    parser = argparse.ArgumentParser(description="Generate Approval Note DOCX")
    parser.add_argument("--output", "-o", required=True, help="Output file path for the DOCX")
    args = parser.parse_args()

    input_text = sys.stdin.read()
    if not input_text.strip():
        sys.stderr.write("Error: Empty JSON input payload\n")
        sys.exit(1)

    try:
        data = json.loads(input_text)
    except Exception as e:
        sys.stderr.write(f"Error parsing JSON input: {e}\n")
        sys.exit(1)

    try:
        resolved_path = create_document(data, args.output)
        # Print output path to stdout as final result
        print(resolved_path)
    except Exception as e:
        sys.stderr.write(f"Error creating DOCX document: {e}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
